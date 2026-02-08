"""
Document Processor - AI-powered document content extraction.

Handles PDF, text, and other document MIME types using Gemini AI.
"""

import asyncio
import base64
import io
import logging
import time
from concurrent.futures import ThreadPoolExecutor
from typing import Optional

from django.conf import settings

from media_curation.ports.ai_port import ContentProcessorPort
from media_curation.ports.storage_port import StoragePort
from media_curation.domain.models import (
    CurationEvent,
    ProcessorResult,
    TenantConfig,
)
from media_curation.domain.exceptions import AIModelError, StorageError


logger = logging.getLogger(__name__)

# Thread pool for async I/O operations
_executor = ThreadPoolExecutor(max_workers=4)


# Extraction prompt for documents
DOCUMENT_EXTRACTION_PROMPT = """
Analyze this document and extract the following:

1. **Main Text Content**: Extract all readable text, preserving paragraph structure.
2. **Key Metadata**: Identify document title, author, date, and any headers/sections.
3. **Structured Data**: Extract any tables, lists, or structured content as JSON.

Format your response as follows:
---TEXT---
[Extracted text content here]
---METADATA---
{
  "title": "...",
  "author": "...",
  "date": "...",
  "sections": ["..."],
  "language": "..."
}
---STRUCTURED---
{
  "tables": [...],
  "lists": [...],
  "key_value_pairs": {...}
}
"""


class DocumentProcessor(ContentProcessorPort):
    """
    Document content processor using Gemini AI.

    Handles PDF, text, and office document MIME types.
    """

    _SUPPORTED_MIME_TYPES = [
        "application/pdf",
        "text/plain",
        "text/html",
        "text/markdown",
        "text/csv",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    def __init__(
        self,
        storage: Optional[StoragePort] = None,
        api_key: Optional[str] = None,
        model_name: Optional[str] = None,
    ):
        """
        Initialize the document processor.

        Args:
            storage: Storage adapter for downloading files
            api_key: Gemini API key (uses settings if not provided)
            model_name: Model to use (uses settings if not provided)
        """
        self.storage = storage

        # Lazy import to avoid issues when google-generativeai is not installed
        try:
            import google.generativeai as genai

            self._genai_available = True
        except ImportError:
            self._genai_available = False
            logger.warning("google-generativeai not installed, using mock mode")
            return

        # Get config from settings
        ai_config = getattr(settings, "MEDIA_CURATION", {}).get("AI_MODEL", {})
        self.api_key = api_key or ai_config.get(
            "API_KEY", getattr(settings, "GOOGLE_API_KEY", None)
        )
        self.model_name = model_name or ai_config.get("MODEL", "gemini-2.0-flash")

        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel(self.model_name)
            logger.info(f"Document processor initialized with model: {self.model_name}")
        else:
            self._genai_available = False
            logger.warning("No GOOGLE_API_KEY found, using mock mode")

    @property
    def supported_mime_types(self) -> list[str]:
        """Return list of MIME types this processor supports."""
        return self._SUPPORTED_MIME_TYPES

    async def process(
        self,
        event: CurationEvent,
        tenant_config: Optional[TenantConfig] = None,
    ) -> ProcessorResult:
        """Process document content and extract text/metadata."""
        start_time = time.time()

        if not self._genai_available:
            # Mock mode: return sample extraction
            return ProcessorResult(
                extracted_text=f"[Mock extraction] Content from {event.raw_gcs_uri}",
                struct_data={
                    "title": "Mock Document",
                    "mime_type": event.mime_type,
                    "mock": True,
                },
                confidence_score=0.0,
                processing_time_ms=int((time.time() - start_time) * 1000),
                language_code="en",
            )

        # Download file content
        try:
            if self.storage:
                content = await self.storage.download_as_bytes(event.raw_gcs_uri)
            else:
                raise StorageError("Storage adapter not configured")
        except Exception as e:
            logger.error(f"Failed to download document: {e}")
            raise StorageError(f"Failed to download document: {e}")

        # Process with Gemini
        def _process():
            import json

            try:
                # Prepare content for Gemini
                if event.mime_type.startswith("text/"):
                    # Text content can be sent directly
                    text_content = content.decode("utf-8", errors="replace")
                    prompt = (
                        f"{DOCUMENT_EXTRACTION_PROMPT}\n\n"
                        f"Document content:\n{text_content[:50000]}"
                    )
                    response = self.model.generate_content(prompt)
                elif event.mime_type == (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ):
                    # .docx files: extract text first since Gemini
                    # doesn't support these as inline_data
                    text_content = self._extract_docx_text(content, event.mime_type)
                    prompt = (
                        f"{DOCUMENT_EXTRACTION_PROMPT}\n\n"
                        f"Document content:\n{text_content[:50000]}"
                    )
                    response = self.model.generate_content(prompt)
                elif event.mime_type == "application/msword":
                    # Legacy .doc files: binary format, cannot be parsed
                    # by python-docx. Attempt basic text extraction.
                    logger.warning(
                        "Legacy .doc format detected; text extraction may "
                        "be incomplete. Consider converting to .docx."
                    )
                    text_content = self._extract_doc_text_fallback(content)
                    prompt = (
                        f"{DOCUMENT_EXTRACTION_PROMPT}\n\n"
                        f"Document content:\n{text_content[:50000]}"
                    )
                    response = self.model.generate_content(prompt)
                else:
                    # Binary content (PDF, images, etc.) sent as inline data
                    encoded_content = base64.b64encode(content).decode("utf-8")
                    response = self.model.generate_content(
                        [
                            DOCUMENT_EXTRACTION_PROMPT,
                            {
                                "inline_data": {
                                    "mime_type": event.mime_type,
                                    "data": encoded_content,
                                }
                            },
                        ]
                    )

                # Parse response
                response_text = response.text
                extracted_text = ""
                struct_data = {}
                metadata = {}

                # Parse sections from response
                if "---TEXT---" in response_text:
                    parts = response_text.split("---TEXT---")
                    if len(parts) > 1:
                        text_part = parts[1]
                        if "---METADATA---" in text_part:
                            extracted_text = text_part.split("---METADATA---")[
                                0
                            ].strip()
                        else:
                            extracted_text = text_part.strip()

                if "---METADATA---" in response_text:
                    parts = response_text.split("---METADATA---")
                    if len(parts) > 1:
                        metadata_part = parts[1]
                        if "---STRUCTURED---" in metadata_part:
                            metadata_part = metadata_part.split("---STRUCTURED---")[0]
                        try:
                            # Find JSON in metadata section
                            json_start = metadata_part.find("{")
                            json_end = metadata_part.rfind("}") + 1
                            if json_start >= 0 and json_end > json_start:
                                metadata = json.loads(
                                    metadata_part[json_start:json_end]
                                )
                        except json.JSONDecodeError:
                            pass

                if "---STRUCTURED---" in response_text:
                    parts = response_text.split("---STRUCTURED---")
                    if len(parts) > 1:
                        struct_part = parts[1].strip()
                        try:
                            json_start = struct_part.find("{")
                            json_end = struct_part.rfind("}") + 1
                            if json_start >= 0 and json_end > json_start:
                                struct_data = json.loads(
                                    struct_part[json_start:json_end]
                                )
                        except json.JSONDecodeError:
                            pass

                # Merge metadata into struct_data
                struct_data["metadata"] = metadata

                # Detect language
                language_code = metadata.get("language", "en")

                return ProcessorResult(
                    extracted_text=extracted_text or response_text[:10000],
                    struct_data=struct_data,
                    confidence_score=0.85,
                    processing_time_ms=int((time.time() - start_time) * 1000),
                    language_code=language_code,
                )

            except Exception as e:
                logger.error(f"Gemini processing error: {e}")
                raise AIModelError(f"Document processing failed: {e}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(_executor, _process)

    @staticmethod
    def _extract_docx_text(content: bytes, mime_type: str) -> str:
        """
        Extract text from .docx files using python-docx.

        Falls back to raw text extraction if python-docx
        is not installed or parsing fails.
        """
        try:
            import docx

            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning(
                "python-docx not installed; falling back to raw text extraction"
            )
        except Exception as e:
            logger.error(f"Failed to parse .docx with python-docx: {e}")

        # Fallback for .docx when python-docx unavailable
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return content.decode("latin-1", errors="replace")

    @staticmethod
    def _extract_doc_text_fallback(content: bytes) -> str:
        """
        Best-effort text extraction from legacy .doc (OLE2) files.

        Legacy .doc is a binary format. This extracts printable ASCII
        runs, which captures body text but loses formatting. For
        higher-fidelity extraction, convert to .docx first.
        """
        import re

        # Extract runs of printable ASCII (>= 4 chars) from the binary
        text_runs = re.findall(rb'[\x20-\x7E]{4,}', content)
        text = "\n".join(run.decode("ascii", errors="replace") for run in text_runs)
        if not text.strip():
            text = "[Unable to extract text from legacy .doc format]"
        return text
